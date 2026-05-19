import os
import math
import streamlit as st
from dotenv import load_dotenv
from groq import Groq
from moviepy.video.io.VideoFileClip import VideoFileClip
from pydub import AudioSegment
from docx import Document
import streamlit.components.v1 as components
from utils.leitores import (
    ler_pdf,
    ler_txt,
    ler_docx
)

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import (
    getSampleStyleSheet
)

load_dotenv()

client = Groq()

#pasta temporária
folder_temp = "temp"

os.makedirs(
    folder_temp,
    exist_ok=True
)

#arquivos temporários
file_audio_temp = (
    f"{folder_temp}/audio.mp3"
)

file_video_temp = (
    f"{folder_temp}/video.mp4"
)


def dividir_audio(
    caminho_audio,
    pasta_saida,
    duracao_minutos=10
):
    """Divide áudio em partes"""

    audio = AudioSegment.from_file(
        caminho_audio
    )

    duracao_ms = (
        duracao_minutos * 60 * 1000
    )

    total_partes = math.ceil(
        len(audio) / duracao_ms
    )

    arquivos = []

    for i in range(total_partes):

        inicio = i * duracao_ms

        fim = inicio + duracao_ms

        parte = audio[inicio:fim]

        caminho_parte = (
            f"{pasta_saida}/parte_{i}.mp3"
        )

        parte.export(
            caminho_parte,
            format="mp3"
        )

        arquivos.append(caminho_parte)

    return arquivos


def transcrever_audio(
    file_audio,
    prompt=None
):
    """Função para transcrever áudio"""

    try:

        if file_audio:

            transcription = (
                client.audio.transcriptions.create(

                    # AJUSTE 1
                    file=(
                        file_audio.name,
                        file_audio
                    ),

                    model="whisper-large-v3",
                    language="pt",
                    response_format="text",
                    prompt=prompt
                )
            )

            return transcription

    except Exception as erro:

        st.error(
            f"Erro no áudio: {erro}"
        )

    return None


def transcrever_video(
    file_video,
    prompt=None
):
    """Função para transcrever vídeos grandes"""

    try:

        if file_video:

            progresso_video = st.progress(0)

            status_video = st.empty()

            # ETAPA 1
            status_video.text(
                "Salvando vídeo..."
            )

            progresso_video.progress(10)

            with open(
                file_video_temp,
                "wb"
            ) as f_video:

                f_video.write(
                    file_video.read()
                )

            # ETAPA 2
            status_video.text(
                "Convertendo vídeo para áudio..."
            )

            progresso_video.progress(30)

            video_convert = (
                VideoFileClip(
                    file_video_temp
                )
            )

            if video_convert.audio is None:

                st.error(
                    "O vídeo não possui áudio."
                )

                return None

            video_convert.audio.write_audiofile(
                file_audio_temp,
                logger=None
            )

            video_convert.close()

            # ETAPA 3
            status_video.text(
                "Dividindo áudio em partes..."
            )

            progresso_video.progress(50)

            partes = dividir_audio(
                file_audio_temp,
                folder_temp,
                duracao_minutos=10
            )

            texto_final = ""

            # barra principal
            progress = st.progress(0)

            status = st.empty()

            #transcrição
            for i, parte in enumerate(partes):

                status.text(
                    f"Transcrevendo parte "
                    f"{i+1} de {len(partes)}..."
                )

                with open(
                    parte,
                    "rb"
                ) as audio:

                    transcricao = (
                        client.audio.transcriptions.create(

                            file=(
                                os.path.basename(parte),
                                audio
                            ),

                            model="whisper-large-v3",
                            language="pt",
                            response_format="text",
                            prompt=prompt
                        )
                    )

                    texto_final += (
                        f"\n\n"
                        f"--- PARTE {i+1} ---\n\n"
                    )

                    texto_final += transcricao

                progresso = int(
                    ((i + 1) / len(partes)) * 100
                )

                progress.progress(
                    progresso
                )

                os.remove(parte)

            progresso_video.progress(100)

            status_video.success(
                "Vídeo processado!"
            )

            status.success(
                "Transcrição finalizada!"
            )

            if os.path.exists(file_audio_temp):

                os.remove(file_audio_temp)

            if os.path.exists(file_video_temp):

                os.remove(file_video_temp)

            return texto_final

    except Exception as erro:

        st.error(
            f"Erro no vídeo: {erro}"
        )

    return None


def gerar_resumo(texto):
    """Gera resumo inteligente"""

    try:

        prompt_resumo = f"""
        Você é um assistente especializado
        em resumir transcrições.

        Analise a transcrição abaixo e gere:

        - resumo geral
        - principais pontos
        - decisões importantes
        - nomes citados
        - datas importantes
        - ações mencionadas

        Transcrição:
        {texto}
        """

        resposta = client.chat.completions.create(

            model="llama-3.3-70b-versatile",

            messages=[
                {
                    "role": "user",
                    "content": prompt_resumo
                }
            ],

            temperature=0.3
        )

        resumo = (
            resposta
            .choices[0]
            .message
            .content
        )

        return resumo

    except Exception as erro:

        st.error(
            f"Erro ao gerar resumo: {erro}"
        )

        return None


def gerar_docx(texto):
    """Gera arquivo DOCX"""

    caminho_docx = (
        f"{folder_temp}/transcricao.docx"
    )

    # AJUSTE 3
    if os.path.exists(caminho_docx):

        os.remove(caminho_docx)

    documento = Document()

    documento.add_heading(
        "Transcrição",
        level=1
    )

    documento.add_paragraph(texto)

    documento.save(caminho_docx)

    return caminho_docx


def gerar_pdf(texto):
    """Gera arquivo PDF"""

    caminho_pdf = (
        f"{folder_temp}/transcricao.pdf"
    )

    # AJUSTE 3
    if os.path.exists(caminho_pdf):

        os.remove(caminho_pdf)

    doc = SimpleDocTemplate(
        caminho_pdf
    )

    styles = getSampleStyleSheet()

    elementos = []

    paragrafos = texto.split("\n")

    for paragrafo in paragrafos:

        elementos.append(
            Paragraph(
                paragrafo,
                styles["BodyText"]
            )
        )

        elementos.append(
            Spacer(1, 12)
        )

    doc.build(elementos)

    return caminho_pdf


def melhorar_texto(texto):
    """Melhora formatação da transcrição"""

    try:

        prompt_formatacao = f"""
        Você é um editor profissional.
        Reescreva a transcrição abaixo
        deixando o texto:

        - organizado
        - com pontuação correta
        - separado em parágrafos
        - agradável de ler
        - mantendo EXATAMENTE o conteúdo
        - sem resumir
        - sem remover informações

        Apenas organize melhor a escrita.

        Transcrição:
        {texto}
        """

        resposta = client.chat.completions.create(

            model="llama-3.3-70b-versatile",

            messages=[
                {
                    "role": "user",
                    "content": prompt_formatacao
                }
            ],

            temperature=0.2
        )

        texto_melhorado = (
            resposta
            .choices[0]
            .message
            .content
        )

        return texto_melhorado

    except Exception as erro:

        st.error(
            f"Erro ao melhorar texto: {erro}"
        )

        return texto


def exibir_transcricao(texto):
    """Mostra transcrição"""

    st.success(
        "Transcrição concluída!"
    )

    # limpa excesso de espaços
    texto_formatado = (
        texto
        .replace(" .", ".")
        .replace(" ,", ",")
        .replace(" ?", "?")
        .replace(" !", "!")
    )

    # add quebra de linha após frases
    texto_formatado = (
        texto_formatado
        .replace(". ", ".\n")
        .replace("? ", "?\n")
        .replace("! ", "!\n")
    )

    # separa partes do vídeo
    texto_formatado = (
        texto_formatado
        .replace("--- PARTE", "\n\n## 🎬 PARTE")
    )

    # transcrição
    st.markdown(
        "## 📄 Transcrição"
    )

    st.markdown(
        f"""
        <div style="
            background-color: #ffffff;
            padding: 25px;
            border-radius: 16px;
            line-height: 1.9;
            font-size: 16px;
            color: #111827;
            white-space: pre-wrap;
            border: 1px solid #e5e7eb;
            box-shadow: 0 2px 8px rgba(0,0,0,0.05);
        ">
        {texto_formatado}
        </div>
        """,
        unsafe_allow_html=True
    )

    # resumo IA
    with st.spinner(
        "Gerando resumo inteligente..."
    ):

        resumo = gerar_resumo(texto)

    if resumo:

        st.markdown(
            "## 🧠 Resumo Inteligente"
        )

        st.markdown(
            f"""
            <div style="
                background-color: #ffffff;
                padding: 25px;
                border-radius: 16px;
                line-height: 1.9;
                font-size: 16px;
                color: #111827;
                border: 1px solid #e5e7eb;
                box-shadow: 0 2px 8px rgba(0,0,0,0.05);
            ">
            {resumo}
            </div>
            """,
            unsafe_allow_html=True
        )

    st.markdown("---")

    st.markdown(
        "## 📥 Downloads"
    )

    # gera arquivos
    caminho_docx = gerar_docx(texto)
    caminho_pdf = gerar_pdf(texto)

    # TXT
    st.download_button(
        label="📄 Baixar TXT",
        data=texto,
        file_name="transcricao.txt",
        mime="text/plain",
        key="download_txt"
    )

    # DOCX
    with open(
        caminho_docx,
        "rb"
    ) as docx_file:

        st.download_button(
            label="📝 Baixar DOCX",
            data=docx_file,
            file_name="transcricao.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            key="download_docx"
        )

    # PDF
    with open(
        caminho_pdf,
        "rb"
    ) as pdf_file:

        st.download_button(
            label="📚 Baixar PDF",
            data=pdf_file,
            file_name="transcricao.pdf",
            mime="application/pdf",
            key="download_pdf"
        )
        
def carregar_estilo():
    """Carrega CSS e HTML personalizado"""

    # CSS
    with open(
        "style.css",
        encoding="utf-8"
    ) as f:

        st.markdown(
            f"<style>{f.read()}</style>",
            unsafe_allow_html=True
        )

    # HEADER HTML
    with open(
        "components/header.html",
        encoding="utf-8"
    ) as f:

        html_code = f.read()

        components.html(
            html_code,
            height=220
        )
        

def main():
    """Função principal"""

    st.set_page_config(
        page_title="App Transcript",
        page_icon="🎙️",
        layout="centered"
    )

    # CARREGA CSS E HTML
    carregar_estilo()

    #abas
    tabs = [
    "Vídeo",
    "Áudio",
    "Arquivos"
    ]

    tab_video, tab_audio, tab_arquivo = st.tabs(
    tabs
    )

    #aba áudio
    with tab_audio:

        st.markdown(
            "### Transcrição de áudio"
        )

        prompt_audio = st.text_input(
            "Prompt do áudio (opcional)",
            key="audio_prompt"
        )

        file_audio = st.file_uploader(
            "Adicione um arquivo de áudio",
            type=[
                "mp3",
                "wav",
                "m4a",
                "ogg"
            ]
        )

        if file_audio:

            st.audio(file_audio)

            with st.spinner(
                "Transcrevendo áudio..."
            ):

                transcricao_audio = (
                    transcrever_audio(
                        file_audio,
                        prompt_audio
                    )
                )

            if transcricao_audio:

                exibir_transcricao(
                    transcricao_audio
                )
    # aba arquivos
    with tab_arquivo:

        st.markdown(
            "### Resumo de documentos"
        )

        arquivo = st.file_uploader(
            "Adicione um arquivo",
            type=[
                "pdf",
                "txt",
                "docx"
            ]
        )

        if arquivo:

            texto = ""

            #PDF
            if arquivo.name.endswith(".pdf"):

                texto = ler_pdf(arquivo)

            #TXT
            elif arquivo.name.endswith(".txt"):

                texto = ler_txt(arquivo)

            #DOCX
            elif arquivo.name.endswith(".docx"):

                texto = ler_docx(arquivo)

            st.markdown(
                "## 📄 Conteúdo Extraído"
            )

            st.text_area(
                "Texto",
                texto,
                height=300
            )

            with st.spinner(
                "Gerando resumo inteligente..."
            ):

                resumo = gerar_resumo(texto)

            if resumo:

                st.markdown(
                    "##🧠 Resumo Inteligente"
                )

                st.markdown(resumo)

    #aba vídeo
    with tab_video:

        st.markdown(
            "### Transcrição de vídeo"
        )

        prompt_video = st.text_input(
            "Prompt do vídeo (opcional)",
            key="video_prompt"
        )

        file_video = st.file_uploader(
            "Adicione um arquivo de vídeo",
            type=[
                "mp4",
                "mov",
                "avi",
                "mkv"
            ]
        )

        if file_video:

            st.video(file_video)

            with st.spinner(
                "Preparando vídeo..."
            ):

                transcricao_video = (
                    transcrever_video(
                        file_video,
                        prompt_video
                    )
                )

            if transcricao_video:

                exibir_transcricao(
                    transcricao_video
                )


if __name__ == "__main__":
    main()


# streamlit run .\transcricao_video.py